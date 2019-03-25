from docx import Document
from docx.shared import Inches

from utils.common import getQuestions, getTriples, getTriplesWithId

questions = getQuestions('data/question.txt')
triples = getTriples('output/output.txt')

doc = Document()  # doc对象
for (i, question) in enumerate(questions):
    string = question
    images = 'images/%s.png' % i  # 保存在本地的图片
    doc.add_paragraph(str(i) + '.' + string)  # 添加文字
    doc.add_picture(images, width=Inches(10))  # 添加图, 设置宽度

    triple_list = getTriplesWithId(i, triples)
    for triple in triple_list:
        doc.add_paragraph(triple)

doc.save('output/output.docx')  # 保存路径
